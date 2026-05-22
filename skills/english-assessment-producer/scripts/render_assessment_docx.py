#!/usr/bin/env python3
"""Render an English assessment DOCX from assessment.json."""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


SKILL_ROOT = Path(__file__).resolve().parents[1]

DEFAULT_FONT = "Times New Roman"
DEFAULT_SIZE = Pt(12)
UNDERLINE_PATTERN = re.compile(r"__([^_\s](?:.*?[^_\s])?)__")
SPEAKER_TURN_SPLIT_PATTERN = re.compile(r"(?=\s*[A-ZÀ-Ỹ][A-Za-zÀ-ỹ' -]{0,24}:)")
LEADING_QUESTION_LABEL_PATTERN = re.compile(
    r"^\s*(?:question|câu)\s*\d+\s*[\.:)]\s*|^\s*[A-Za-z]?\d+\s*[\.:)]\s*",
    re.IGNORECASE,
)
LEADING_SECTION_LABEL_PATTERN = re.compile(r"^\s*section\s+[A-Za-z0-9]+:\s*", re.IGNORECASE)


def apply_normal_style_typography(document: Document) -> None:
    """Set Times New Roman 12pt as document defaults, matching format_rules."""
    style = document.styles["Normal"]
    style.font.name = DEFAULT_FONT
    style.font.size = DEFAULT_SIZE
    style.font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    for style_name, size, bold in [
        ("Title", 16, True),
        ("Heading 1", 12, True),
        ("Heading 2", 12, True),
        ("Heading 3", 12, True),
    ]:
        if style_name in document.styles:
            heading = document.styles[style_name]
            heading.font.name = DEFAULT_FONT
            heading.font.size = Pt(size)
            heading.font.bold = bold
            heading.font.color.rgb = RGBColor(0, 0, 0)
            heading.paragraph_format.line_spacing = 1.0
            heading.paragraph_format.space_before = Pt(0 if style_name == "Title" else 4)
            heading.paragraph_format.space_after = Pt(2)


def set_run_font(run, size: Pt | None = None, bold: bool = False) -> None:
    run.font.name = DEFAULT_FONT
    run.font.size = size or DEFAULT_SIZE
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bool(bold)


def add_markup_runs(paragraph, text: str, bold: bool = False) -> None:
    """Add text runs and render __underlined text__ as Word underline."""
    position = 0
    for match in UNDERLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, bold=bold)
        run = paragraph.add_run(match.group(1))
        set_run_font(run, bold=bold)
        run.underline = True
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, bold=bold)


def remove_markup(text: str) -> str:
    return UNDERLINE_PATTERN.sub(r"\1", text)


def add_plain_runs(paragraph, text: str, bold: bool = False) -> None:
    run = paragraph.add_run(remove_markup(text))
    set_run_font(run, bold=bold)


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 2, line: float = 1.0) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_hanging_indent(paragraph, indent: float = 14, first: float = -14) -> None:
    pf = paragraph.paragraph_format
    pf.left_indent = Mm(indent * 0.176)  # pt to mm approx
    pf.first_line_indent = Mm(first * 0.176)


@dataclass
class RenderState:
    numbering: str = "per_section"
    document_profile: str = "worksheet"
    option_layout: str = "auto"
    global_counter: int = 0
    container_counters: dict[str, int] = field(default_factory=dict)


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, dict):
        raise ValueError("Assessment source must be a JSON object.")
    return data


def safe_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def option_label(index: int) -> str:
    return chr(ord("A") + index)


def add_paragraph(document: Document, text: str = "", style: str | None = None, bold: bool = False):
    try:
        p = document.add_paragraph(text, style=style)
    except KeyError:
        p = document.add_paragraph(text)
    if text and p.runs:
        set_run_font(p.runs[0], bold=bold)
    return p


def add_heading(document: Document, text: str, level: int = 1):
    try:
        p = document.add_heading(text, level=level)
    except KeyError:
        p = document.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(16 if level == 1 else 13)
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p
    # Normalize heading font to Times New Roman
    size = Pt(16 if level == 0 else 12)
    for run in p.runs:
        set_run_font(run, size=size, bold=True)
    set_paragraph_spacing(p, before=0 if level == 0 else 4, after=2)
    return p


def build_document(template: Path | None = None) -> Document:
    if template and template.exists():
        return Document(str(template))
    doc = Document()
    apply_normal_style_typography(doc)
    return doc


def set_cell_text(cell, text: str, bold: bool = False, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    set_paragraph_spacing(paragraph, before=0, after=0)
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def default_page_setup(profile: str) -> dict[str, Any]:
    if profile == "exam":
        return {
            "paper": "A4",
            "margins_mm": {"top": 10, "bottom": 11, "left": 20, "right": 10},
        }
    return {
        "paper": "worksheet_practice",
        "width_mm": 211.67,
        "height_mm": 273.93,
        "margins_mm": {"top": 7.5, "bottom": 5, "left": 7.7, "right": 6.5},
    }


def apply_page_setup(document: Document, data: dict[str, Any]) -> None:
    metadata = data.get("metadata") if isinstance(data.get("metadata"), dict) else {}
    profile = metadata.get("document_profile") or metadata.get("document_type") or "worksheet"
    rendering = data.get("rendering") if isinstance(data.get("rendering"), dict) else {}
    setup = rendering.get("page_setup") if isinstance(rendering.get("page_setup"), dict) else default_page_setup(profile)
    for section in document.sections:
        if setup.get("paper") == "A4":
            section.page_width = Mm(210)
            section.page_height = Mm(297)
        elif setup.get("paper") == "worksheet_practice":
            section.page_width = Mm(float(setup.get("width_mm", 211.67)))
            section.page_height = Mm(float(setup.get("height_mm", 273.93)))
        margins = setup.get("margins_mm")
        if isinstance(margins, dict):
            if margins.get("top") is not None:
                section.top_margin = Mm(float(margins["top"]))
            if margins.get("bottom") is not None:
                section.bottom_margin = Mm(float(margins["bottom"]))
            if margins.get("left") is not None:
                section.left_margin = Mm(float(margins["left"]))
            if margins.get("right") is not None:
                section.right_margin = Mm(float(margins["right"]))


def write_exam_header(document: Document, header: dict[str, Any], metadata: dict[str, Any]) -> None:
    left_lines = header.get("left") or []
    right_lines = header.get("right") or []
    if not left_lines and not right_lines:
        return

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    left_cell, right_cell = table.rows[0].cells
    set_cell_text(left_cell, "\n".join(safe_text(x) for x in left_lines), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(right_cell, "\n".join(safe_text(x) for x in right_lines), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    note = metadata.get("paper_note")
    if note:
        p = add_paragraph(document, safe_text(note))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=2)


def write_footer(document: Document, metadata: dict[str, Any], rendering: dict[str, Any]) -> None:
    footer_config = rendering.get("footer") if isinstance(rendering.get("footer"), dict) else {}
    if footer_config.get("enabled") is False:
        return
    exam_code = footer_config.get("exam_code") or metadata.get("exam_code")
    if not exam_code and not footer_config.get("page_numbers"):
        return

    for section in document.sections:
        paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        paragraph.text = ""
        if exam_code:
            set_run_font(paragraph.add_run(f"Mã đề {exam_code}\tPage "))
        elif footer_config.get("page_numbers"):
            set_run_font(paragraph.add_run("Page "))
        if footer_config.get("page_numbers", True):
            add_page_field(paragraph, "PAGE")
            set_run_font(paragraph.add_run("/"))
            add_page_field(paragraph, "NUMPAGES")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(paragraph, before=0, after=0)


def write_metadata(document: Document, data: dict[str, Any]) -> None:
    metadata = data.get("metadata") if isinstance(data.get("metadata"), dict) else {}
    rendering = data.get("rendering") if isinstance(data.get("rendering"), dict) else {}
    header = data.get("header") if isinstance(data.get("header"), dict) else {}

    if metadata.get("document_profile") == "exam" or header:
        write_exam_header(document, header, metadata)
    else:
        title = safe_text(metadata.get("title") or "English Assessment")
        heading = add_heading(document, title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_parts = []
    if metadata.get("document_type"):
        info_parts.append(f"Type: {metadata['document_type']}")
    if metadata.get("grade"):
        info_parts.append(f"Grade: {metadata['grade']}")
    if metadata.get("duration_minutes"):
        info_parts.append(f"Time: {metadata['duration_minutes']} minutes")
    if metadata.get("total_points"):
        info_parts.append(f"Total points: {metadata['total_points']}")
    if info_parts and metadata.get("document_profile") != "exam":
        paragraph = add_paragraph(document, " | ".join(info_parts))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, before=0, after=2)

    units = metadata.get("units") or []
    if units and metadata.get("document_profile") != "exam":
        unit_text = "; ".join(
            f"Unit {u.get('unit_number')}: {u.get('title')}" if isinstance(u, dict) else safe_text(u)
            for u in units
        )
        set_paragraph_spacing(add_paragraph(document, f"Scope: {unit_text}"), before=0, after=2)

    student_fields = metadata.get("student_fields") or ["Name", "Class"]
    if metadata.get("document_profile") == "exam":
        table = document.add_table(rows=1, cols=max(2, len(student_fields) + (1 if metadata.get("exam_code") else 0)))
        table.autofit = True
        cells = table.rows[0].cells
        for idx, field in enumerate(student_fields):
            set_cell_text(cells[idx], f"{field}: " + "." * 40)
        if metadata.get("exam_code"):
            set_cell_text(cells[-1], f"Mã đề {metadata['exam_code']}")
    elif student_fields:
        set_paragraph_spacing(
            add_paragraph(document, "    ".join(f"{field}: ____________________" for field in student_fields)),
            before=0,
            after=2,
        )

    write_footer(document, metadata, rendering)


def write_matrix(document: Document, matrix: list[Any]) -> None:
    rows = [row for row in matrix if isinstance(row, dict)]
    if not rows:
        return

    add_heading(document, "Test Matrix", level=1)
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Section", "Knowledge", "Exercise", "Difficulty", "Count", "Points"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], safe_text(row.get("title") or row.get("section_id") or row.get("block_id")))
        set_cell_text(cells[1], safe_text(row.get("knowledge_type")))
        set_cell_text(cells[2], safe_text(row.get("exercise_type")))
        set_cell_text(cells[3], safe_text(row.get("difficulty")))
        set_cell_text(cells[4], safe_text(row.get("count")))
        points = row.get("total_points")
        if points is None and row.get("count") is not None and row.get("points_each") is not None:
            try:
                points = float(row["count"]) * float(row["points_each"])
            except (TypeError, ValueError):
                points = ""
        set_cell_text(cells[5], safe_text(points))


def write_generic_table(document: Document, columns: list[Any], rows: list[Any], header: bool = True) -> None:
    if not columns or not rows:
        return
    table = document.add_table(rows=1 if header else 0, cols=len(columns))
    table.style = "Table Grid"
    table.autofit = True
    if header:
        for idx, column in enumerate(columns):
            set_cell_text(table.rows[0].cells[idx], safe_text(column), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        values = row if isinstance(row, list) else [row.get(c, "") for c in columns] if isinstance(row, dict) else [row]
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            set_cell_text(cell, safe_text(values[idx]) if idx < len(values) else "")


def write_word_bank(document: Document, block: dict[str, Any]) -> None:
    words = block.get("words") or block.get("items") or []
    if not words:
        return
    columns = int(block.get("columns_count") or min(4, max(1, len(words))))
    rows = [words[i : i + columns] for i in range(0, len(words), columns)]
    write_generic_table(document, ["" for _ in range(columns)], rows, header=False)


def write_crossword(document: Document, block: dict[str, Any]) -> None:
    grid = block.get("grid") or []
    if not grid:
        return
    max_cols = max(len(row) for row in grid if isinstance(row, list))
    table = document.add_table(rows=0, cols=max_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in grid:
        cells = table.add_row().cells
        for idx in range(max_cols):
            set_cell_text(cells[idx], safe_text(row[idx]) if idx < len(row) else "", align=WD_ALIGN_PARAGRAPH.CENTER)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def write_notice_box(document: Document, block: dict[str, Any]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(table.rows[0].cells[0], safe_text(block.get("text")), bold=bool(block.get("bold")), align=WD_ALIGN_PARAGRAPH.CENTER)


def write_text_block(document: Document, block: dict[str, Any]) -> None:
    text = safe_text(block.get("text"))
    if not text:
        return
    if block.get("style") == "heading":
        add_heading(document, text, level=int(block.get("level") or 1))
    else:
        for paragraph_text in re.split(r"\n\s*\n", text.strip()):
            if paragraph_text:
                add_paragraph(document, paragraph_text)


def next_question_number(state: RenderState, container_id: str) -> int:
    if state.numbering == "global":
        state.global_counter += 1
        return state.global_counter
    state.container_counters[container_id] = state.container_counters.get(container_id, 0) + 1
    return state.container_counters[container_id]


def question_label(state: RenderState, container_id: str, question: dict[str, Any]) -> str:
    if question.get("display_label"):
        return safe_text(question["display_label"])
    if state.numbering == "global" or state.document_profile == "exam":
        return f"Question {next_question_number(state, container_id)}"
    return safe_text(next_question_number(state, container_id))


def normalize_for_compare(text: str) -> str:
    text = re.sub(r"[^a-z0-9]+", " ", text.lower())
    return re.sub(r"\s+", " ", text).strip()


def strip_leading_question_label(text: str) -> str:
    text = LEADING_SECTION_LABEL_PATTERN.sub("", text)
    return LEADING_QUESTION_LABEL_PATTERN.sub("", text).strip()


def clean_display_label(label: str) -> str:
    return re.sub(r"[\s\.:)]+$", "", safe_text(label).strip())


def answer_blank_suffix(stem: str, exercise_type: str | None) -> str:
    if exercise_type != "reading_tf":
        return ""
    if re.search(r"_{3,}|\.{4,}|\bT\s*/\s*F\s*/\s*NG\b", stem, flags=re.IGNORECASE):
        return ""
    return " ________"


def append_given_word_if_needed(stem: str, question: dict[str, Any], exercise_type: str | None) -> str:
    if exercise_type != "rewrite_with_given_word":
        return stem
    given_word = safe_text(question.get("given_word") or question.get("cue_word")).strip()
    if not given_word:
        return stem
    if given_word.lower() in stem.lower():
        return stem
    return f"{stem} ({given_word})" if stem else f"({given_word})"


def rewrite_prompt_text(question: dict[str, Any], exercise_type: str | None) -> str:
    if exercise_type not in {"sentence_rewrite", "rewrite_with_given_word"}:
        return ""
    prompt = safe_text(question.get("prompt")).strip()
    if prompt:
        return prompt
    start = safe_text(
        question.get("answer_start") or question.get("rewrite_start") or question.get("response_start")
    ).strip()
    if start:
        return f"=> {start} {'_' * 58}"
    return f"=> {'_' * 62}"


def should_render_answer_lines(question: dict[str, Any], exercise_type: str | None, response_prompt: str) -> bool:
    if exercise_type == "question_making":
        return True
    if not question.get("lines"):
        return False
    if exercise_type in {"sentence_rewrite", "rewrite_with_given_word"} and re.search(r"_{3,}", response_prompt):
        return False
    return True


def is_redundant_question_stem(stem: str, group: dict[str, Any]) -> bool:
    cleaned_stem = strip_leading_question_label(stem)
    normalized_stem = normalize_for_compare(cleaned_stem)
    if not normalized_stem:
        return True
    for key in ["instructions", "title"]:
        normalized_group_text = normalize_for_compare(safe_text(group.get(key)))
        if normalized_group_text and (
            normalized_stem == normalized_group_text
            or normalized_stem in normalized_group_text
            or normalized_group_text in normalized_stem
        ):
            return True
    exercise_type = group.get("exercise_type")
    generic_stems = {
        "pronunciation_odd_one": [
            "choose the word whose underlined part is pronounced differently from the others",
            "choose the word that has a different sound pattern from the others",
        ],
        "stress_odd_one": [
            "choose the word that has a different stress pattern from the others",
            "choose the word whose stress pattern is different from the others",
        ],
    }
    return normalized_stem in {normalize_for_compare(x) for x in generic_stems.get(exercise_type, [])}


def write_question_label_and_stem(document: Document, state: RenderState, label: str, stem: str, suffix: str):
    p = add_paragraph(document)
    label = clean_display_label(label)
    if state.document_profile == "exam":
        label_run = p.add_run(f"{label}.")
        set_run_font(label_run, bold=True)
        if stem or suffix:
            set_run_font(p.add_run(" "))
            add_markup_runs(p, f"{stem}{suffix}")
    else:
        set_run_font(p.add_run(f"{label}."))
        if stem or suffix:
            set_run_font(p.add_run(" "))
            add_markup_runs(p, f"{stem}{suffix}")
        set_hanging_indent(p, indent=14, first=-14)
    set_paragraph_spacing(p, before=0, after=0)
    return p


def normalize_options(options: Any) -> list[tuple[str, str]]:
    if isinstance(options, dict):
        return [(safe_text(k), safe_text(v)) for k, v in options.items()]
    if isinstance(options, list):
        result = []
        for idx, value in enumerate(options):
            if isinstance(value, dict):
                label = safe_text(value.get("label") or option_label(idx))
                text = safe_text(value.get("text") or value.get("value"))
            else:
                label = option_label(idx)
                text = safe_text(value)
            text = re.sub(rf"^\s*{re.escape(label)}\s*[\.:)]\s*", "", text, flags=re.IGNORECASE).strip()
            result.append((label, text))
        return result
    return []


def option_text_is_long(normalized: list[tuple[str, str]]) -> bool:
    return any(len(remove_markup(text)) > 24 for _, text in normalized)


def write_option_run(paragraph, label: str, text: str, trailing: str = "", allow_markup: bool = True) -> None:
    label_run = paragraph.add_run(f"{label}. ")
    set_run_font(label_run, bold=True)
    if allow_markup:
        add_markup_runs(paragraph, text)
    else:
        add_plain_runs(paragraph, text)
    if trailing:
        set_run_font(paragraph.add_run(trailing))


def write_options_table(
    document: Document,
    normalized: list[tuple[str, str]],
    allow_markup: bool = True,
    leading_label: str | None = None,
    two_rows: bool = False,
) -> None:
    if two_rows:
        column_count = 3 if leading_label else 2
        table = document.add_table(rows=0, cols=column_count)
        table.autofit = False
        pairs = [normalized[0:2], normalized[2:4]]
        for row_index, pair in enumerate(pairs):
            cells = table.add_row().cells
            option_offset = 1 if leading_label else 0
            if leading_label:
                cells[0].width = Mm(8)
                p_label = cells[0].paragraphs[0]
                p_label.clear()
                if row_index == 0:
                    set_run_font(p_label.add_run(f"{clean_display_label(leading_label)}."))
                set_paragraph_spacing(p_label, before=0, after=0)
            for idx, (label, text) in enumerate(pair):
                cells[idx + option_offset].width = Mm(82 if leading_label else 90)
                p = cells[idx + option_offset].paragraphs[0]
                p.clear()
                write_option_run(p, label, text, allow_markup=allow_markup)
                set_paragraph_spacing(p, before=0, after=0)
        return

    column_count = len(normalized) + (1 if leading_label else 0)
    table = document.add_table(rows=1, cols=column_count)
    table.autofit = False
    cells = table.rows[0].cells
    option_offset = 1 if leading_label else 0
    if leading_label:
        cells[0].width = Mm(8)
        p_label = cells[0].paragraphs[0]
        p_label.clear()
        set_run_font(p_label.add_run(f"{clean_display_label(leading_label)}."))
        set_paragraph_spacing(p_label, before=0, after=0)
    for idx, (label, text) in enumerate(normalized):
        cells[idx + option_offset].width = Mm(42 if leading_label else 47)
        p = cells[idx + option_offset].paragraphs[0]
        p.clear()
        write_option_run(p, label, text, allow_markup=allow_markup)
        set_paragraph_spacing(p, before=0, after=0)


def write_options(document: Document, options: Any, layout: str, allow_markup: bool = True) -> None:
    normalized = normalize_options(options)
    if not normalized:
        return
    if layout == "stacked" or (layout == "auto" and any(len(remove_markup(text)) > 90 for _, text in normalized)):
        for label, text in normalized:
            p = add_paragraph(document)
            write_option_run(p, label, text, allow_markup=allow_markup)
            set_paragraph_spacing(p, before=0, after=0)
        return
    if layout == "two_column" or (layout == "auto" and option_text_is_long(normalized)):
        write_options_table(document, normalized, allow_markup=allow_markup, two_rows=True)
        return
    write_options_table(document, normalized, allow_markup=allow_markup, two_rows=False)


def write_labeled_options(document: Document, state: RenderState, label: str, options: Any, allow_markup: bool = True) -> None:
    normalized = normalize_options(options)
    if not normalized:
        return
    write_options_table(
        document,
        normalized,
        allow_markup=allow_markup,
        leading_label=label,
        two_rows=option_text_is_long(normalized),
    )


def write_items_table(document: Document, items: Any) -> None:
    if not isinstance(items, list) or not items:
        return
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "A", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.rows[0].cells[1], "B", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for item in items:
        cells = table.add_row().cells
        if isinstance(item, dict):
            left_text = safe_text(item.get("left") or item.get("term") or item.get("prompt"))
            right_text = safe_text(item.get("right") or item.get("definition") or item.get("match"))
        else:
            left_text = safe_text(item)
            right_text = ""
        for idx, text in enumerate([left_text, right_text]):
            cell = cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = DEFAULT_FONT
            run.font.size = DEFAULT_SIZE
            run.font.color.rgb = RGBColor(0, 0, 0)


def write_question(document: Document, state: RenderState, container_id: str, question: dict[str, Any], group: dict[str, Any]) -> None:
    if question.get("passage"):
        p = add_paragraph(document, safe_text(question["passage"]))
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    label = question_label(state, container_id, question)
    q_type = question.get("exercise_type") or group.get("exercise_type")
    raw_stem = question.get("stem")
    if raw_stem is None and q_type not in {"sentence_rewrite", "rewrite_with_given_word"}:
        raw_stem = question.get("prompt")
    stem = strip_leading_question_label(safe_text(raw_stem or ""))
    stem = append_given_word_if_needed(stem, question, q_type)
    if is_redundant_question_stem(stem, group):
        stem = ""
    points = question.get("points")
    suffix = answer_blank_suffix(stem, q_type)
    if points is not None:
        suffix += f" ({points} point{'s' if points != 1 else ''})"
    layout = question.get("option_layout") or group.get("option_layout") or state.option_layout
    if not stem and not suffix and question.get("options") and q_type in {"pronunciation_odd_one", "stress_odd_one", "odd_one_topic"}:
        write_labeled_options(document, state, label, question.get("options"), allow_markup=q_type != "stress_odd_one")
        return
    if q_type == "reading_gap_fill" and question.get("options"):
        write_labeled_options(document, state, label, question.get("options"), allow_markup=False)
        return
    write_question_label_and_stem(document, state, label, safe_text(stem), suffix)
    response_prompt = rewrite_prompt_text(question, q_type)
    if response_prompt:
        p = add_paragraph(document)
        add_markup_runs(p, response_prompt)
        set_hanging_indent(p, indent=14, first=0)
        set_paragraph_spacing(p, before=0, after=0)
    if q_type != "reading_tf":
        write_options(document, question.get("options"), layout, allow_markup=q_type != "stress_odd_one")
    write_items_table(document, question.get("items"))
    if should_render_answer_lines(question, q_type, response_prompt):
        line_count = int(question.get("lines") or 1)
        for _ in range(line_count):
            p = add_paragraph(document, "________________________________________________________________")
            for run in p.runs:
                run.font.name = DEFAULT_FONT
                run.font.size = DEFAULT_SIZE
                run.font.color.rgb = RGBColor(0, 0, 0)


def write_question_group(document: Document, block: dict[str, Any], state: RenderState) -> None:
    block_id = safe_text(block.get("id") or "group")
    title_text = safe_text(block.get("title"))
    if block.get("title"):
        if state.document_profile == "exam":
            paragraph = add_paragraph(document, title_text)
            for run in paragraph.runs:
                run.bold = True
                run.font.name = DEFAULT_FONT
                run.font.size = DEFAULT_SIZE
                run.font.color.rgb = RGBColor(0, 0, 0)
            set_paragraph_spacing(paragraph, before=3, after=2)
        else:
            add_heading(document, title_text, level=int(block.get("level") or 1))
    if block.get("instructions"):
        instruction_text = safe_text(block["instructions"])
        if normalize_for_compare(instruction_text) != normalize_for_compare(title_text):
            p = add_paragraph(document)
            add_markup_runs(p, instruction_text)
            p.paragraph_format.space_after = Pt(4)
    if block.get("passage"):
        p = add_paragraph(document, safe_text(block["passage"]))
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = DEFAULT_FONT
            run.font.size = DEFAULT_SIZE
            run.font.color.rgb = RGBColor(0, 0, 0)
    if block.get("notice"):
        write_notice_box(document, {"text": block.get("notice"), "bold": block.get("notice_bold", True)})
    for question in block.get("questions") or []:
        if isinstance(question, dict):
            write_question(document, state, block_id, question, block)


def write_block(document: Document, block: dict[str, Any], state: RenderState) -> None:
    block_type = block.get("type")
    if block.get("title") and block_type not in {"question_group", "heading"}:
        add_heading(document, safe_text(block["title"]), level=int(block.get("level") or 1))
    if block_type == "heading":
        add_heading(document, safe_text(block.get("text") or block.get("title")), level=int(block.get("level") or 1))
    elif block_type == "text":
        write_text_block(document, block)
    elif block_type in {"vocabulary_table", "resource_table"}:
        write_generic_table(document, block.get("columns") or [], block.get("rows") or [], header=True)
    elif block_type == "word_bank":
        write_word_bank(document, block)
    elif block_type == "crossword":
        write_crossword(document, block)
    elif block_type == "notice_box":
        write_notice_box(document, block)
    elif block_type == "passage":
        write_text_block(document, {"text": block.get("text") or block.get("passage")})
    elif block_type == "question_group":
        write_question_group(document, block, state)


def write_sections_as_blocks(document: Document, sections: list[Any], state: RenderState) -> None:
    for section in sections:
        if not isinstance(section, dict):
            continue
        block = dict(section)
        block["type"] = "question_group"
        write_question_group(document, block, state)


def collect_answers_from_questions(data: dict[str, Any], state: RenderState | None = None) -> list[tuple[str, str]]:
    answers: list[tuple[str, str]] = []

    def add_questions(container: dict[str, Any], container_id: str):
        local_state = state or RenderState(
            numbering=(data.get("rendering") or {}).get("question_numbering", "per_section"),
            document_profile=(data.get("metadata") or {}).get("document_profile", "worksheet"),
            option_layout=(data.get("rendering") or {}).get("option_layout", "auto"),
        )
        for question in container.get("questions") or []:
            if isinstance(question, dict) and "answer" in question:
                label = question.get("display_label") or question.get("id")
                if not label and state:
                    label = question_label(local_state, container_id, question)
                answers.append((safe_text(label), safe_text(question.get("answer"))))

    if isinstance(data.get("blocks"), list):
        for idx, block in enumerate(data["blocks"], start=1):
            if isinstance(block, dict) and block.get("type") == "question_group":
                add_questions(block, safe_text(block.get("id") or idx))
    else:
        for idx, section in enumerate(data.get("sections") or [], start=1):
            if isinstance(section, dict):
                add_questions(section, safe_text(section.get("id") or idx))
    return answers


def collect_answers(data: dict[str, Any]) -> list[tuple[str, str]]:
    raw_key = data.get("answer_key")
    if isinstance(raw_key, list) and raw_key:
        answers = []
        for item in raw_key:
            if isinstance(item, dict):
                answers.append((safe_text(item.get("question_id")), safe_text(item.get("answer"))))
            else:
                answers.append(("", safe_text(item)))
        return answers
    if isinstance(raw_key, dict) and raw_key:
        return [(safe_text(k), safe_text(v)) for k, v in raw_key.items()]
    return collect_answers_from_questions(data)


def write_answer_key(document: Document, data: dict[str, Any]) -> bool:
    rendering = data.get("rendering") if isinstance(data.get("rendering"), dict) else {}
    if rendering.get("include_answer_key") is False:
        return False
    answers = collect_answers(data)
    if not answers:
        return False
    document.add_page_break()
    add_heading(document, "Answer Key", level=1)
    for question_id, answer in answers:
        label = f"{question_id}: " if question_id else ""
        p = add_paragraph(document, f"{label}{answer}")
        for run in p.runs:
            run.font.name = DEFAULT_FONT
            run.font.size = DEFAULT_SIZE
            run.font.color.rgb = RGBColor(0, 0, 0)
    return True


def transcript_lines(transcript: str) -> list[str]:
    lines: list[str] = []
    for raw_line in transcript.strip().splitlines():
        line = raw_line.strip()
        if not line:
            lines.append("")
            continue
        parts = SPEAKER_TURN_SPLIT_PATTERN.split(line)
        split_parts = [part.strip() for part in parts if part.strip()]
        lines.extend(split_parts or [line])
    return lines


def write_transcript(document: Document, data: dict[str, Any], answer_key_written: bool = False) -> None:
    rendering = data.get("rendering") if isinstance(data.get("rendering"), dict) else {}
    if rendering.get("include_transcript") is False:
        return
    listening = data.get("listening")
    if not isinstance(listening, dict) or not listening.get("transcript"):
        return
    if not answer_key_written:
        document.add_page_break()
    add_heading(document, "Listening Transcript", level=1)
    transcript = safe_text(listening.get("transcript"))
    for line in transcript_lines(transcript):
        if line:
            add_paragraph(document, line)
        else:
            add_paragraph(document, "")


def normalize_paragraph(paragraph) -> None:
    if paragraph.paragraph_format.space_before is None:
        paragraph.paragraph_format.space_before = Pt(0)
    if paragraph.paragraph_format.space_after is None:
        paragraph.paragraph_format.space_after = Pt(2)
    if paragraph.paragraph_format.line_spacing is None:
        paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        if run.text:
            run.font.name = DEFAULT_FONT
            run.font.color.rgb = RGBColor(0, 0, 0)
            if run.font.size is None:
                run.font.size = DEFAULT_SIZE


def normalize_document(document: Document) -> None:
    apply_normal_style_typography(document)
    for paragraph in document.paragraphs:
        normalize_paragraph(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    normalize_paragraph(paragraph)
    for section in document.sections:
        for part in [section.header, section.footer]:
            for paragraph in part.paragraphs:
                normalize_paragraph(paragraph)


def render(input_path: Path, output_path: Path, template: Path | None = None) -> None:
    data = load_json(input_path)
    document = build_document(template)
    apply_page_setup(document, data)
    metadata = data.get("metadata") if isinstance(data.get("metadata"), dict) else {}
    rendering = data.get("rendering") if isinstance(data.get("rendering"), dict) else {}
    profile = metadata.get("document_profile") or metadata.get("document_type") or "worksheet"
    state = RenderState(
        numbering=rendering.get("question_numbering") or ("global" if profile == "exam" else "per_section"),
        document_profile=profile,
        option_layout=rendering.get("option_layout") or "auto",
    )

    write_metadata(document, data)
    print_matrix = rendering.get("print_matrix")
    if print_matrix is True or (print_matrix is None and data.get("matrix") and not data.get("blocks")):
        write_matrix(document, data.get("matrix") if isinstance(data.get("matrix"), list) else [])

    if isinstance(data.get("blocks"), list):
        for block in data["blocks"]:
            if isinstance(block, dict):
                write_block(document, block, state)
    else:
        write_sections_as_blocks(document, data.get("sections") if isinstance(data.get("sections"), list) else [], state)

    answer_key_written = write_answer_key(document, data)
    write_transcript(document, data, answer_key_written=answer_key_written)
    normalize_document(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="Path to assessment.json")
    parser.add_argument("--output", required=True, type=Path, help="DOCX output path")
    parser.add_argument("--template", type=Path, default=None, help="Optional DOCX template override. Default uses built-in format_rules.")
    parser.add_argument(
        "--allow-blank-template",
        action="store_true",
        help="Deprecated; blank documents with built-in format_rules are now the default.",
    )
    args = parser.parse_args()

    try:
        render(args.input, args.output, args.template)
    except Exception as exc:  # noqa: BLE001 - CLI should report concise render failures.
        print(f"Failed to render DOCX: {exc}", file=sys.stderr)
        return 1

    print(f"Rendered DOCX: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
