#!/usr/bin/env python3
"""
KB Extraction Script — Global Success THCS
Extracts vocabulary, grammar, lesson structure from giáo án .docx files
→ inserts into SQLite knowledge base (knowledge_base.db)

Usage:
    python3 kb_extract.py                  # extract all grades
    python3 kb_extract.py --grade 6        # extract grade 6 only
    python3 kb_extract.py --dry-run        # print without inserting

Schema overview:
    units            — unit metadata (grade, unit_number, title, topic)
    lessons          — lesson list per unit (7 lessons per unit)
    vocabulary       — words from Language Analysis tables (IPA, meaning EN/VI)
    grammar_points   — grammar rules from grammar tables (form, examples)
    topic_contexts   — reading/listening topic hints (no full content needed)
    questions        — question bank generated later (empty at extraction time)
    matrices         — test matrix templates (populated manually or via skill)
"""

import os
import re
import json
import sqlite3
import argparse
from pathlib import Path
from docx import Document

# ─── Config ──────────────────────────────────────────────────────────────────

BASE_DIR = Path(__file__).parent
DB_PATH = BASE_DIR / "knowledge_base.db"

GRADE_FOLDERS = {
    6: "GA TA 6",
    7: "GA TA 7",
    8: "GA TA 8",
    9: "GA TA 9",
}

# Lesson name → skill focus mapping (Global Success structure)
LESSON_SKILL_MAP = {
    "getting started":      "overview",
    "a closer look 1":      "vocabulary,pronunciation",
    "a closer look 2":      "grammar",
    "communication":        "communication",
    "skills 1":             "reading,speaking",
    "skills 2":             "listening,writing",
    "looking back":         "review",
    "project":              "review",
    "review":               "review",
}

# ─── DB Init ─────────────────────────────────────────────────────────────────

SCHEMA = """
CREATE TABLE IF NOT EXISTS units (
    id           INTEGER PRIMARY KEY AUTOINCREMENT,
    grade        INTEGER NOT NULL,
    unit_number  INTEGER NOT NULL,
    title        TEXT,
    topic        TEXT,
    UNIQUE(grade, unit_number)
);

CREATE TABLE IF NOT EXISTS lessons (
    id             INTEGER PRIMARY KEY AUTOINCREMENT,
    unit_id        INTEGER NOT NULL REFERENCES units(id) ON DELETE CASCADE,
    lesson_number  INTEGER,
    lesson_name    TEXT,
    skill_focus    TEXT,     -- comma-separated: vocabulary,pronunciation / grammar / reading,speaking / etc.
    UNIQUE(unit_id, lesson_number)
);

CREATE TABLE IF NOT EXISTS vocabulary (
    id               INTEGER PRIMARY KEY AUTOINCREMENT,
    unit_id          INTEGER NOT NULL REFERENCES units(id) ON DELETE CASCADE,
    lesson_id        INTEGER REFERENCES lessons(id),
    word             TEXT NOT NULL,
    part_of_speech   TEXT,             -- n / v / adj / adv / phrase
    pronunciation_ipa TEXT,
    meaning_en       TEXT,
    meaning_vi       TEXT,
    UNIQUE(unit_id, word)
);

CREATE TABLE IF NOT EXISTS grammar_points (
    id               INTEGER PRIMARY KEY AUTOINCREMENT,
    unit_id          INTEGER NOT NULL REFERENCES units(id) ON DELETE CASCADE,
    lesson_id        INTEGER REFERENCES lessons(id),
    point_name       TEXT NOT NULL,    -- e.g. "Present Simple", "Adverbs of frequency"
    form_description TEXT,             -- e.g. "S + V-inf/V(s/es) + ..."
    examples         TEXT              -- JSON array of example strings
);

CREATE TABLE IF NOT EXISTS topic_contexts (
    id               INTEGER PRIMARY KEY AUTOINCREMENT,
    unit_id          INTEGER NOT NULL REFERENCES units(id) ON DELETE CASCADE,
    lesson_id        INTEGER REFERENCES lessons(id),
    skill            TEXT NOT NULL,    -- 'reading' | 'listening'
    topic_hint       TEXT              -- short description of topic/context
);

CREATE TABLE IF NOT EXISTS questions (
    id               INTEGER PRIMARY KEY AUTOINCREMENT,
    unit_id          INTEGER NOT NULL REFERENCES units(id) ON DELETE CASCADE,
    lesson_id        INTEGER REFERENCES lessons(id),
    knowledge_type   TEXT NOT NULL,    -- vocabulary | grammar | reading | listening
    exercise_type    TEXT NOT NULL,    -- MCQ | TF | fill_blank | matching | ordering | writing | translation
    difficulty       INTEGER DEFAULT 1 CHECK(difficulty IN (1,2,3)),  -- 1=easy 2=medium 3=hard
    content          TEXT,             -- JSON: question stem, options, passage, etc.
    answer           TEXT,             -- JSON: correct answer(s)
    used_count       INTEGER DEFAULT 0,
    last_used_date   TEXT,
    created_at       TEXT DEFAULT (datetime('now'))
);

CREATE TABLE IF NOT EXISTS matrices (
    id          INTEGER PRIMARY KEY AUTOINCREMENT,
    name        TEXT NOT NULL,
    grade       INTEGER,              -- NULL = applies to all grades
    description TEXT,
    structure   TEXT NOT NULL         -- JSON array of sections
    -- structure example:
    -- [
    --   {"section": "Vocabulary", "knowledge_type": "vocabulary",
    --    "exercise_type": "MCQ", "count": 5, "difficulty": 1},
    --   {"section": "Grammar",    "knowledge_type": "grammar",
    --    "exercise_type": "MCQ", "count": 5, "difficulty": 2},
    --   {"section": "Reading",    "knowledge_type": "reading",
    --    "exercise_type": "TF",  "count": 3, "difficulty": 2},
    --   {"section": "Listening",  "knowledge_type": "listening",
    --    "exercise_type": "MCQ", "count": 4, "difficulty": 2}
    -- ]
);

CREATE INDEX IF NOT EXISTS idx_vocab_unit     ON vocabulary(unit_id);
CREATE INDEX IF NOT EXISTS idx_vocab_word     ON vocabulary(word);
CREATE INDEX IF NOT EXISTS idx_grammar_unit   ON grammar_points(unit_id);
CREATE INDEX IF NOT EXISTS idx_questions_unit ON questions(unit_id);
CREATE INDEX IF NOT EXISTS idx_questions_type ON questions(exercise_type, knowledge_type);
CREATE INDEX IF NOT EXISTS idx_questions_used ON questions(used_count);
"""

def init_db(conn: sqlite3.Connection):
    conn.execute("PRAGMA journal_mode = DELETE")  # virtiofs-safe (no WAL)
    conn.executescript(SCHEMA)
    conn.commit()

# ─── Parsing Helpers ─────────────────────────────────────────────────────────

def parse_unit_number(filename: str) -> int | None:
    """Extract unit number from filename. Handles: Unit 1.docx, TA7-Unit 1.docx"""
    m = re.search(r'[Uu]nit\s+(\d+)', filename)
    return int(m.group(1)) if m else None

def is_review_file(filename: str) -> bool:
    return 'review' in filename.lower() or 'pl ' in filename.lower()

def get_table_header(table) -> list[str]:
    """Return lowercased cell texts of first row."""
    if not table.rows:
        return []
    return [c.text.strip().lower() for c in table.rows[0].cells]

def is_vocab_table(table) -> bool:
    """Detect Language Analysis vocabulary table.
    Header contains 'pronunciation' or 'ipa' AND ('meaning' or 'vietnamese').
    """
    header = ' '.join(get_table_header(table))
    return (
        ('pronunciation' in header or '/ipa' in header or 'ipa' in header)
        and ('meaning' in header or 'vietnamese' in header)
        and 'teacher' not in header
    )

def is_grammar_table(table) -> bool:
    """Detect grammar/language rule table.

    Handles patterns found across all 4 grades:
      P1 direct:  row0 has (form|structure|grammar) + (example|usage)
      P2 titled:  row0 = repeated grammar-point name; row1 has P1 keywords
      P3 column:  row0 has example/examples/usage in any cell (no form keyword needed)
                  — covers: "| positive | comparative | rule",
                             "if-clause | main clause | note",
                             "types of sentences | definition | examples",
                             "making requests | examples", etc.

    Excludes:
      - Vocabulary tables (have 'pronunciation' + 'vietnamese')
      - Activity tables (have 'teacher'/'student'/'task'/'stage'/'anticipated')
      - Data-only tables (school info, fact tables)
    """
    if not table.rows:
        return False

    _EXCLUDE = ('pronunciation', 'vietnamese', 'teacher', 'student',
                'task', 'stage', 'anticipated', 'solution', 'school',
                'time | activity', 'detailed plan', 'what are the jobs',
                'no of english', 'family member', 'my future city')
    _FORM_KW = ('form', 'structure', 'grammar', 'if-clause', 'direct speech',
                'reported speech', 'language function', 'how we use',
                'types of', 'relative pronoun', 'adverbial')
    _EXAMPLE_KW = ('example', 'examples', 'usage', 'rule', 'note', 'definition')

    def _cell_texts(row) -> list[str]:
        return [c.text.strip().lower() for c in row.cells]

    def _joined(cells: list[str]) -> str:
        return ' '.join(cells)

    def _is_excluded(joined: str) -> bool:
        return any(ex in joined for ex in _EXCLUDE)

    def _has_grammar_signal(cells: list[str]) -> bool:
        j = _joined(cells)
        has_form = any(kw in j for kw in _FORM_KW)
        has_example = any(kw in j for kw in _EXAMPLE_KW)
        return (has_form or has_example) and not _is_excluded(j)

    def _is_title_row(cells: list[str]) -> bool:
        """Row where all non-empty cells have the same text (merged title pattern).
        Check single-cell length, not joined — title may repeat across merged cols.
        """
        non_empty = [c for c in cells if c.strip()]
        if not non_empty:
            return False
        first = non_empty[0]
        # Single cell too long = probably a paragraph, not a title
        if len(first) > 120:
            return False
        if _is_excluded(first.lower()):
            return False
        # All cells same text (merged) OR single cell
        return len(set(non_empty)) == 1

    row0_cells = _cell_texts(table.rows[0])
    row0_j = _joined(row0_cells)

    # P1: direct — row0 itself is a grammar header
    if _has_grammar_signal(row0_cells):
        return True

    if len(table.rows) >= 2:
        row1_cells = _cell_texts(table.rows[1])

        # P2: titled — row0 is grammar-point name, row1 is grammar header
        if _is_title_row(row0_cells) and _has_grammar_signal(row1_cells):
            return True

        # P3: column pattern — row0 has example/rule/note/definition anywhere
        # even when first cell is empty or has a structural label
        row0_has_example = any(kw in row0_j for kw in _EXAMPLE_KW)
        if row0_has_example and not _is_excluded(row0_j):
            return True

    return False

def parse_part_of_speech(word_cell: str) -> tuple[str, str]:
    """
    Parse word and part of speech from cell like '1. subject (n)' or 'boarding school (n)'.
    Returns (word_clean, pos).
    """
    # Remove leading number + dot: "1. subject (n)" → "subject (n)"
    text = re.sub(r'^\d+\.\s*', '', word_cell.strip())
    # Extract POS from parentheses at end
    pos_match = re.search(r'\(([^)]+)\)\s*$', text)
    pos = pos_match.group(1).strip() if pos_match else ''
    word = re.sub(r'\s*\([^)]+\)\s*$', '', text).strip()
    return word, pos

def extract_ipa(text: str) -> str:
    """Extract IPA from cell — content between /.../ marks."""
    m = re.search(r'(/[^/]+/)', text)
    return m.group(1).strip() if m else text.strip()

def extract_vocab_rows(table) -> list[dict]:
    """Extract vocabulary entries from a Language Analysis table."""
    rows = []
    for row in table.rows[1:]:  # skip header
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue
        word_cell = cells[0]
        if not word_cell or word_cell.lower() in ('form', ''):
            continue
        word, pos = parse_part_of_speech(word_cell)
        if not word:
            continue

        # Handle merged cells — cols may be: Form | Pronunciation | Meaning | Vietnamese
        # or compressed 3-col versions
        ipa = ''
        meaning_en = ''
        meaning_vi = ''

        if len(cells) >= 4:
            ipa = extract_ipa(cells[1])
            meaning_en = cells[2].strip()
            meaning_vi = cells[3].strip()
        elif len(cells) == 3:
            ipa = extract_ipa(cells[1])
            meaning_en = cells[2].strip()

        rows.append({
            'word': word,
            'part_of_speech': pos,
            'pronunciation_ipa': ipa,
            'meaning_en': meaning_en,
            'meaning_vi': meaning_vi,
        })
    return rows

def extract_grammar_rows(table) -> list[dict]:
    """Extract grammar point entries from a grammar/language table.

    Returns list of {point_name, form_description, examples}.
    Handles all patterns:
      - Direct (Form|Example, Structure|Examples, Grammar|Usage, etc.)
      - Titled (grammar name row + header row + data rows)
      - Multi-column (| Positive | Comparative | Rule, Types|Definition|Examples, etc.)
    """
    if not table.rows:
        return []

    rows = table.rows
    point_name = ''
    data_start = 1

    row0_cells = [c.text.strip() for c in rows[0].cells]
    row0_joined = ' | '.join(row0_cells).lower()

    _HEADER_KW = ('form', 'structure', 'grammar', 'example', 'examples',
                  'usage', 'rule', 'definition', 'if-clause', 'language function',
                  'how we use', 'types of', 'direct speech', 'reported speech',
                  'positive', 'comparative', 'note', 'relative pronoun',
                  'adverbial', 'making', 'asking', 'expressing', 'giving',
                  'persuading', 'respond', 'ask for', 'express')
    _SKIP_CONTENT = ('teacher', 'student', 'anticipated', 'pronunciation',
                     'vietnamese', 'school', 'task')

    def _is_header(cells: list[str]) -> bool:
        j = ' '.join(cells).lower()
        return any(kw in j for kw in _HEADER_KW) and not any(s in j for s in _SKIP_CONTENT)

    def _is_title(cells: list[str]) -> bool:
        non_empty = [c for c in cells if c.strip()]
        return bool(non_empty) and len(set(non_empty)) == 1 and len(non_empty[0]) < 80

    # Determine structure: titled vs direct
    if _is_title(row0_cells) and len(rows) >= 2:
        point_name = row0_cells[0].strip()
        data_start = 2  # skip title + header
        # Validate row1 is actually a header
        row1_cells = [c.text.strip() for c in rows[1].cells]
        if not _is_header(row1_cells):
            data_start = 1  # no header after title
    elif _is_header(row0_cells):
        point_name = ''
        data_start = 1
    else:
        data_start = 0  # treat all rows as content

    # Collect content: flatten each row into form_part | ... | example
    content_rows = []
    for row in rows[data_start:]:
        cells = [c.text.strip() for c in row.cells]
        # Skip empty rows or rows that are sub-headers
        row_text = ' '.join(cells)
        if not row_text.strip():
            continue
        if all(not c for c in cells):
            continue
        content_rows.append(cells)

    if not content_rows:
        return []

    # Represent content: first col = form/structure, rest = examples/notes
    forms = []
    examples = []
    for cells in content_rows:
        if len(cells) >= 1 and cells[0]:
            forms.append(cells[0])
        if len(cells) >= 2:
            # Collect all non-empty cells from col 1 onward as examples
            ex_parts = [c for c in cells[1:] if c]
            if ex_parts:
                examples.append(' | '.join(ex_parts))

    if not forms and not examples:
        return []

    return [{'point_name': point_name,
             'form_description': ' || '.join(forms),
             'examples': examples}]

def detect_lesson_sections(doc) -> list[dict]:
    """
    Walk paragraphs and detect lesson headings.
    Returns list of {lesson_number, lesson_name, skill_focus, para_index}.
    """
    sections = []
    lesson_pattern = re.compile(
        r'[Ll]esson\s+(\d+)\s*[:\-–]\s*(.*)', re.IGNORECASE
    )
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        m = lesson_pattern.match(text)
        if m:
            num = int(m.group(1))
            name_raw = m.group(2).strip()
            # Clean subtitle after dash/em-dash
            name_clean = re.split(r'\s*[–\-]\s*', name_raw)[0].strip()
            skill = _lesson_skill(name_clean)
            sections.append({
                'lesson_number': num,
                'lesson_name': name_clean or name_raw,
                'skill_focus': skill,
                'para_index': i,
            })
    return sections

def extract_grammar_from_paragraphs(doc, lesson_sections: list[dict]) -> list[dict]:
    """
    Fallback: extract grammar from paragraph text when no grammar table found.
    Handles 3 patterns observed across missing units:

    P1 — Language analysis block (G9U1, G9U11):
        After "Language analysis" paragraph, before "Assumption" or "III."
        Lines starting with –, +, or bare structure text.

    P2 — OBJECTIVES Knowledge block (G7U11):
        Between "1. Knowledge" and "2. Competences/Core competence"
        Lines describing grammar rules (contain Eg:, +, a., b., will/V-ing etc.)

    P3 — Activity description (G6U11):
        "Grammar point presentation: X" → extract X as point name only.

    Returns list of {point_name, form_description, examples, lesson_number}.
    """
    paras = doc.paragraphs
    results = []

    # Find Lesson 3 paragraph index
    lesson3_idx = None
    lesson4_idx = len(paras)
    for sec in lesson_sections:
        if sec['lesson_number'] == 3:
            lesson3_idx = sec['para_index']
        if sec['lesson_number'] == 4:
            lesson4_idx = sec['para_index']

    if lesson3_idx is None:
        return []

    l3_paras = [p.text.strip() for p in paras[lesson3_idx:lesson4_idx]]

    # ── P1: Language analysis block ──
    lang_results = _extract_lang_analysis_block(l3_paras)
    if lang_results:
        results.extend(lang_results)

    # ── P2: OBJECTIVES Knowledge block ──
    if not results:
        obj_results = _extract_objectives_grammar(l3_paras)
        if obj_results:
            results.extend(obj_results)

    # ── P3: Grammar point presentation names ──
    if not results:
        name_results = _extract_grammar_point_names(l3_paras)
        results.extend(name_results)

    for r in results:
        r['lesson_number'] = 3
    return results


def _extract_lang_analysis_block(paras: list[str]) -> list[dict]:
    """Extract grammar from Language analysis...Assumption block."""
    in_block = False
    content_lines = []
    for line in paras:
        if line.lower() == 'language analysis':
            in_block = True
            continue
        if in_block:
            if line.lower() in ('assumption', 'assumptions', 'iii. procedures', 'board plan'):
                break
            if line:
                content_lines.append(line)

    if not content_lines:
        return []

    # First non-empty line may be an explanation, rest are structure lines
    forms = []
    examples = []
    for line in content_lines:
        stripped = line.lstrip('–-+•* \t')
        if stripped.startswith('Eg:') or stripped.startswith('E.g') or 'e.g' in stripped.lower()[:10]:
            examples.append(stripped)
        elif stripped:
            forms.append(stripped)

    if not forms:
        return []

    # Grammar point name: infer from first form line
    point_name = forms[0][:60] if forms else 'Grammar point'
    return [{'point_name': point_name,
             'form_description': ' || '.join(forms),
             'examples': examples}]


def _extract_objectives_grammar(paras: list[str]) -> list[dict]:
    """Extract grammar from '1. Knowledge' block in OBJECTIVES."""
    in_knowledge = False
    content_lines = []
    for line in paras:
        ll = line.lower()
        if re.match(r'1\.\s*knowledge', ll):
            in_knowledge = True
            continue
        if in_knowledge:
            if re.match(r'2\.\s*(competenc|core)', ll):
                break
            if line:
                content_lines.append(line)

    if not content_lines:
        return []

    # Skip if block is vocabulary-focused, not grammar
    first_lower = content_lines[0].lower()
    if first_lower.startswith('vocabulary') or first_lower.startswith('- vocabulary'):
        return []

    # Filter for lines with grammar signal (structure patterns, Eg:, a./b.)
    _STRUCT_KW = ('will', 'v-ing', 'v-inf', 'bare inf', 'eg:', 'e.g',
                  's +', 'be +', 'have +', 'had +', 'shall', 'would',
                  'pronoun', 'gerund', 'participle', 'clause',
                  r'^\+', r'^a\.', r'^b\.', r'^eg')
    grammar_lines = []
    for line in content_lines:
        ll = line.lower()
        if any(kw in ll for kw in _STRUCT_KW) or re.match(r'^[a-z]\.\s', ll):
            grammar_lines.append(line)

    if not grammar_lines:
        # Fall back: use all content lines as grammar description
        grammar_lines = content_lines

    # First content line as point name
    point_name = content_lines[0][:80] if content_lines else 'Grammar point'
    forms = []
    examples = []
    for line in grammar_lines:
        stripped = line.lstrip('–-+•* \t')
        if 'eg:' in stripped.lower()[:10] or 'e.g' in stripped.lower()[:10]:
            examples.append(stripped)
        else:
            forms.append(stripped)

    return [{'point_name': point_name,
             'form_description': ' || '.join(forms),
             'examples': examples}]


def _extract_grammar_point_names(paras: list[str]) -> list[dict]:
    """Extract grammar from 'Grammar point presentation: X' lines."""
    results = []
    pattern = re.compile(r'grammar point presentation[:\s]+(.+)', re.IGNORECASE)
    seen = set()
    for line in paras:
        m = pattern.search(line)
        if m:
            name = m.group(1).strip().rstrip('.')
            if name not in seen:
                seen.add(name)
                results.append({'point_name': name,
                                'form_description': '',
                                'examples': []})
    return results


def _lesson_skill(name: str) -> str:
    name_lower = name.lower()
    for key, skill in LESSON_SKILL_MAP.items():
        if key in name_lower:
            return skill
    return 'general'

def extract_unit_title(doc) -> tuple[str, str]:
    """
    Return (unit_title, topic) from first UNIT heading.
    e.g. "UNIT 1: MY NEW SCHOOL" → title="MY NEW SCHOOL", topic=same
    """
    unit_pattern = re.compile(r'UNIT\s+\d+\s*[:\-–]\s*(.*)', re.IGNORECASE)
    for para in doc.paragraphs[:30]:
        m = unit_pattern.match(para.text.strip())
        if m:
            title = m.group(1).strip()
            return title, title
    return '', ''

def detect_skill_from_context(doc, lesson_sections: list[dict]) -> list[dict]:
    """
    Detect reading/listening topic hints from task descriptions near Skills 1/2 lessons.
    Looks for paragraphs mentioning listening/reading topics.
    """
    contexts = []
    para_texts = [p.text.strip() for p in doc.paragraphs]

    for sec in lesson_sections:
        skill = sec.get('skill_focus', '')
        if 'listening' in skill:
            # Look for description near this lesson
            start = sec['para_index']
            snippet = ' '.join(para_texts[start:start+30])
            hint = _extract_topic_hint(snippet, 'listening')
            if hint:
                contexts.append({
                    'lesson_number': sec['lesson_number'],
                    'skill': 'listening',
                    'topic_hint': hint,
                })
        if 'reading' in skill:
            start = sec['para_index']
            snippet = ' '.join(para_texts[start:start+30])
            hint = _extract_topic_hint(snippet, 'reading')
            if hint:
                contexts.append({
                    'lesson_number': sec['lesson_number'],
                    'skill': 'reading',
                    'topic_hint': hint,
                })
    return contexts

def _extract_topic_hint(snippet: str, skill: str) -> str:
    """Extract a short topic description from a text snippet."""
    # Look for "talk about X", "listen to X", "read about X", "topic: X"
    patterns = [
        r'(?:talk|write|speak|learn)\s+about\s+([\w\s,]+?)(?:\.|;|,|\n)',
        r'(?:listen(?:ing)?|read(?:ing)?)\s+(?:about|to|a|the)\s+([\w\s]+?)(?:\.|;|,|\n)',
        r'topic[:\s]+([\w\s]+?)(?:\.|;|,|\n)',
        r'passage[s]?\s+about\s+([\w\s]+?)(?:\.|;|,|\n)',
    ]
    for pat in patterns:
        m = re.search(pat, snippet, re.IGNORECASE)
        if m:
            return m.group(1).strip()[:120]
    return ''

# ─── DB Insert Helpers ────────────────────────────────────────────────────────

def upsert_unit(conn, grade: int, unit_number: int, title: str, topic: str) -> int:
    conn.execute(
        """INSERT INTO units (grade, unit_number, title, topic)
           VALUES (?, ?, ?, ?)
           ON CONFLICT(grade, unit_number) DO UPDATE SET
               title=excluded.title, topic=excluded.topic""",
        (grade, unit_number, title, topic)
    )
    row = conn.execute(
        "SELECT id FROM units WHERE grade=? AND unit_number=?",
        (grade, unit_number)
    ).fetchone()
    return row[0]

def upsert_lesson(conn, unit_id: int, lesson_number: int,
                   lesson_name: str, skill_focus: str) -> int:
    conn.execute(
        """INSERT INTO lessons (unit_id, lesson_number, lesson_name, skill_focus)
           VALUES (?, ?, ?, ?)
           ON CONFLICT(unit_id, lesson_number) DO UPDATE SET
               lesson_name=excluded.lesson_name, skill_focus=excluded.skill_focus""",
        (unit_id, lesson_number, lesson_name, skill_focus)
    )
    row = conn.execute(
        "SELECT id FROM lessons WHERE unit_id=? AND lesson_number=?",
        (unit_id, lesson_number)
    ).fetchone()
    return row[0]

def insert_vocab(conn, unit_id: int, lesson_id: int | None, vocab: dict):
    try:
        conn.execute(
            """INSERT INTO vocabulary
               (unit_id, lesson_id, word, part_of_speech,
                pronunciation_ipa, meaning_en, meaning_vi)
               VALUES (?,?,?,?,?,?,?)
               ON CONFLICT(unit_id, word) DO UPDATE SET
                   lesson_id=excluded.lesson_id,
                   part_of_speech=excluded.part_of_speech,
                   pronunciation_ipa=excluded.pronunciation_ipa,
                   meaning_en=excluded.meaning_en,
                   meaning_vi=excluded.meaning_vi""",
            (unit_id, lesson_id, vocab['word'], vocab['part_of_speech'],
             vocab['pronunciation_ipa'], vocab['meaning_en'], vocab['meaning_vi'])
        )
    except Exception as e:
        print(f"    [WARN] vocab insert fail '{vocab['word']}': {e}")

def insert_grammar(conn, unit_id: int, lesson_id: int | None,
                   point_name: str, entry: dict):
    conn.execute(
        """INSERT INTO grammar_points
           (unit_id, lesson_id, point_name, form_description, examples)
           VALUES (?,?,?,?,?)""",
        (unit_id, lesson_id, point_name,
         entry['form_description'], json.dumps(entry['examples']))
    )

def insert_topic(conn, unit_id: int, lesson_id: int | None,
                 skill: str, hint: str):
    conn.execute(
        """INSERT INTO topic_contexts (unit_id, lesson_id, skill, topic_hint)
           VALUES (?,?,?,?)""",
        (unit_id, lesson_id, skill, hint)
    )

# ─── Per-file Extraction ─────────────────────────────────────────────────────

def extract_file(conn, filepath: Path, grade: int, dry_run: bool = False):
    filename = filepath.name
    unit_number = parse_unit_number(filename)
    if unit_number is None:
        return  # skip Review/PL files

    print(f"  Grade {grade} Unit {unit_number}: {filename}")

    try:
        doc = Document(str(filepath))
    except Exception as e:
        print(f"    [ERROR] Cannot open: {e}")
        return

    # --- Unit metadata
    title, topic = extract_unit_title(doc)
    if not dry_run:
        unit_id = upsert_unit(conn, grade, unit_number, title, topic)
    else:
        unit_id = -1
        print(f"    Unit title: '{title}'")

    # --- Lesson sections
    lessons = detect_lesson_sections(doc)
    lesson_id_map: dict[int, int] = {}  # lesson_number → db id

    for les in lessons:
        if not dry_run:
            lid = upsert_lesson(conn, unit_id, les['lesson_number'],
                                les['lesson_name'], les['skill_focus'])
            lesson_id_map[les['lesson_number']] = lid
        else:
            print(f"    Lesson {les['lesson_number']}: {les['lesson_name']} [{les['skill_focus']}]")

    # --- Tables: iterate doc body in order to track lesson context
    # We'll match each table to nearest preceding lesson heading
    para_to_lesson: dict[int, int] = {}
    lesson_boundaries = sorted(lessons, key=lambda x: x['para_index'])

    # Build paragraph-index → lesson_number lookup
    def get_lesson_for_para(pidx: int) -> int | None:
        current = None
        for sec in lesson_boundaries:
            if sec['para_index'] <= pidx:
                current = sec['lesson_number']
            else:
                break
        return current

    # Walk docx body XML to get paragraph + table order
    from docx.oxml.ns import qn
    body = doc.element.body
    para_idx = 0
    table_idx = 0

    doc_paras = doc.paragraphs
    doc_tables = doc.tables

    vocab_count = 0
    grammar_count = 0

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            # Advance paragraph index
            para_idx += 1

        elif tag == 'tbl':
            if table_idx >= len(doc_tables):
                table_idx += 1
                continue
            table = doc_tables[table_idx]
            table_idx += 1

            lesson_num = get_lesson_for_para(para_idx)
            lesson_id = lesson_id_map.get(lesson_num) if not dry_run else None

            if is_vocab_table(table):
                vocab_rows = extract_vocab_rows(table)
                for v in vocab_rows:
                    if not dry_run:
                        insert_vocab(conn, unit_id, lesson_id, v)
                    else:
                        print(f"    [VOCAB L{lesson_num}] {v['word']} {v['pronunciation_ipa']} — {v['meaning_vi']}")
                vocab_count += len(vocab_rows)

            elif is_grammar_table(table):
                gram_rows = extract_grammar_rows(table)
                for entry in gram_rows:
                    if not entry['form_description']:
                        continue
                    # Use name from table if present, else infer from paragraphs
                    point_name = (entry.get('point_name')
                                  or _infer_grammar_name(doc_paras, para_idx))
                    if not dry_run:
                        insert_grammar(conn, unit_id, lesson_id, point_name, entry)
                    else:
                        print(f"    [GRAMMAR L{lesson_num}] {point_name}: {entry['form_description'][:60]}")
                    grammar_count += 1

    # --- Grammar paragraph fallback (units with no grammar tables)
    if grammar_count == 0:
        para_grammar = extract_grammar_from_paragraphs(doc, lessons)
        for pg in para_grammar:
            lesson_id = lesson_id_map.get(pg.get('lesson_number', 3)) if not dry_run else None
            if not dry_run:
                insert_grammar(conn, unit_id, lesson_id,
                               pg['point_name'],
                               {'form_description': pg['form_description'],
                                'examples': pg['examples']})
            else:
                print(f"    [GRAMMAR-PARA L3] {pg['point_name'][:60]}: {pg['form_description'][:60]}")
            grammar_count += 1

    # --- Topic contexts
    topic_hints = detect_skill_from_context(doc, lessons)
    for th in topic_hints:
        lesson_id = lesson_id_map.get(th['lesson_number']) if not dry_run else None
        if th['topic_hint']:
            if not dry_run:
                insert_topic(conn, unit_id, lesson_id, th['skill'], th['topic_hint'])
            else:
                print(f"    [TOPIC] {th['skill']}: {th['topic_hint']}")

    if not dry_run:
        conn.commit()

    print(f"    → {vocab_count} vocab, {grammar_count} grammar blocks, {len(lessons)} lessons, {len(topic_hints)} topics")


def _infer_grammar_name(paragraphs, near_idx: int) -> str:
    """Look backwards from near_idx for a grammar point name paragraph."""
    # Search up to 15 paragraphs back
    start = max(0, near_idx - 15)
    for para in reversed(paragraphs[start:near_idx]):
        text = para.text.strip()
        # Likely a grammar heading if short, no punctuation, not a task line
        if text and len(text) < 80 and not text.startswith('-') and not text.startswith('Task'):
            # Filter out teacher activity text
            if 'teacher' not in text.lower() and 'student' not in text.lower():
                return text
    return 'Unknown grammar point'

# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Extract KB from giáo án .docx files')
    parser.add_argument('--grade', type=int, choices=[6,7,8,9],
                        help='Extract specific grade only')
    parser.add_argument('--dry-run', action='store_true',
                        help='Print extracted data without inserting to DB')
    parser.add_argument('--db', type=str, default=str(DB_PATH),
                        help=f'SQLite DB path (default: {DB_PATH})')
    args = parser.parse_args()

    grades = [args.grade] if args.grade else [6, 7, 8, 9]

    conn = sqlite3.connect(args.db)
    conn.execute("PRAGMA foreign_keys = ON")
    if not args.dry_run:
        init_db(conn)
        print(f"DB initialized: {args.db}")

    for grade in grades:
        folder_name = GRADE_FOLDERS[grade]
        folder_path = BASE_DIR / folder_name
        if not folder_path.exists():
            print(f"[SKIP] Folder not found: {folder_path}")
            continue

        print(f"\n=== Grade {grade} ({folder_name}) ===")
        docx_files = sorted(folder_path.glob("*.docx"))

        for filepath in docx_files:
            if is_review_file(filepath.name):
                print(f"  [SKIP] {filepath.name}")
                continue
            extract_file(conn, filepath, grade, dry_run=args.dry_run)

    conn.close()

    if not args.dry_run:
        # Print summary
        conn2 = sqlite3.connect(args.db)
        rows = conn2.execute("""
            SELECT u.grade, COUNT(DISTINCT u.id) as units,
                   COUNT(DISTINCT v.id) as vocab,
                   COUNT(DISTINCT g.id) as grammar
            FROM units u
            LEFT JOIN vocabulary v ON v.unit_id = u.id
            LEFT JOIN grammar_points g ON g.unit_id = u.id
            GROUP BY u.grade ORDER BY u.grade
        """).fetchall()
        print("\n─── Extraction Summary ───")
        print(f"{'Grade':<8} {'Units':<8} {'Vocab':<10} {'Grammar':<10}")
        for r in rows:
            print(f"{r[0]:<8} {r[1]:<8} {r[2]:<10} {r[3]:<10}")
        conn2.close()

if __name__ == '__main__':
    main()
